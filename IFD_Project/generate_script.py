import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    doc = Document()
    
    # Configure default styles for extreme readability
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(16)

    # Title
    title = doc.add_heading('Kopitalk Presentation Script', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph('Designed for ultimate readability.')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph() # Spacing

    def add_slide(slide_num, title_text, speaker_lines):
        # Slide Header
        p = doc.add_paragraph()
        run = p.add_run(f"SLIDE {slide_num}: {title_text.upper()}")
        run.bold = True
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0, 112, 192) # Blue
        
        # Add visual separator
        doc.add_paragraph("_" * 40)

        for line in speaker_lines:
            if "CUE:" in line:
                cp = doc.add_paragraph()
                cr = cp.add_run(f"[{line}]")
                cr.bold = True
                cr.italic = True
                cr.font.color.rgb = RGBColor(220, 20, 60) # Red
            else:
                parts = line.split(":", 1)
                if len(parts) == 2:
                    speaker, text = parts
                    sp = doc.add_paragraph()
                    sp.paragraph_format.space_after = Pt(14)
                    
                    # Speaker Name
                    sr = sp.add_run(f"{speaker.strip().upper()}:\n")
                    sr.bold = True
                    sr.font.size = Pt(18)
                    
                    # Change speaker color
                    if "HAZIQ" in speaker.upper():
                        sr.font.color.rgb = RGBColor(0, 153, 0) # Green for Haziq
                    elif "IRFAN" in speaker.upper():
                        sr.font.color.rgb = RGBColor(153, 0, 153) # Purple for Irfan
                    elif "ANDRE" in speaker.upper() or "ADAM" in speaker.upper():
                        sr.font.color.rgb = RGBColor(0, 0, 0) # Black for storytellers

                    # Speaker Text
                    tr = sp.add_run(text.strip())
                    tr.font.size = Pt(16)
        
        doc.add_page_break()

    # SLIDE 1
    add_slide(1, "Title Slide", [
        "CUE: Wait for audience to settle. Press Right Arrow to load presentation.",
        "HAZIQ: Good morning everyone! Thank you for coming down to our showcase. We are incredibly excited to share our final project with you today.",
        "HAZIQ: Our project is called Kopitalk. It's a concept that means a lot to our team, and it's all about bridging connections."
    ])

    # SLIDE 2
    add_slide(2, "What is Kopitalk?", [
        "CUE: Press Right Arrow to bring up slide elements.",
        "ANDRE: Thanks, Haziq. So, what exactly is Kopitalk? Simply put, it is not just a standard video game. It's a hybrid experience.",
        "ANDRE: Kopitalk combines physical board game elements with a smart digital interface. Let me tell you a little bit about the story behind why we built this."
    ])

    # SLIDE 3
    add_slide(3, "The Problem", [
        "CUE: Press Right Arrow",
        "ADAM: The story started when we looked closely at our own families. Generation Z and Alpha are completely glued to screens and digital devices.",
        "ADAM: But our grandparents and older relatives prefer tangible, physical interactions. Because of this, there's a massive generational and technological gap that is actively pulling families apart at the dinner table."
    ])

    # SLIDE 4
    add_slide(4, "The Solution", [
        "CUE: Press Right Arrow",
        "ANDRE: Exactly. Our solution to this problem was to create something that naturally appeals to both sides.",
        "ANDRE: We designed a physical medium that older folks understand and enjoy touching, strictly connected to a digital platform that young people find engaging.",
        "ANDRE: It's a cooperative game designed specifically to force interaction and communication across those boundaries."
    ])

    # SLIDE 5
    add_slide(5, "Unique Aspects", [
        "CUE: Press Right Arrow",
        "ADAM: What makes Kopitalk unique is that we aren't just putting an iPad on a table. We are taking physical objects—real pieces that you can touch and hold—and having them directly affect a virtual world.",
        "ADAM: It's a highly tactile experience that you just can't replicate with a standard phone or video game."
    ])

    # SLIDE 6
    add_slide(6, "The Process (Dashboard)", [
        "CUE: Press Right Arrow to reveal Virtual Dashboard. Press multiple times to animate through the 4 stages.",
        "IRFAN: Thanks Adam. So, how did we actually build this? Getting here took a lot of work.",
        "IRFAN: Our process involved four massive stages: Extensive 'Ideation' to find the perfect mechanics, '3D Design' to shape the physical components, heavy 'Hardware' engineering to make the electronics work, and finally complex 'Software' development to tie the physical world to the digital screen."
    ])

    # SLIDE 7
    add_slide(7, "Pitch Video", [
        "CUE: Laptop will appear. Pitch video is muted and idle.",
        "HAZIQ: Before we dive into the physical prototype itself, we want to show you a quick video that perfectly captures the spirit of what Kopitalk is all about. Take a look.",
        "CUE: Press Right Arrow to zoom video into FULLSCREEN and Un-mute.",
        "CUE: (Wait 30 seconds for video to finish).",
        "CUE: Press Right Arrow to zoom back out into the laptop."
    ])

    # SLIDE 8
    add_slide(8, "Prototype Gallery", [
        "CUE: Press Right Arrow to enter the Coverflow Carousel.",
        "HAZIQ: As you can see, the vision is big! Now, let's talk about the actual physical prototype we built. Irfan and I spent countless hours 3D printing and assembling these components in the lab.",
        "CUE: Press Right Arrow periodically to click through the 3D Prototype photos.",
        "IRFAN: That's right, Haziq. What you're looking at here includes our Modular Structural Frame and the actual 3D printed MRT Station scale model.",
        "IRFAN: We had to do dozens of test prints and assembly tests just to ensure the physical tiles would interface correctly with our grid layout."
    ])

    # SLIDE 9
    add_slide(9, "Demo & Hardware Engine", [
        "CUE: Press Right Arrow.",
        "IRFAN: The hardware is driven by a completely custom setup underneath the board. Every time a piece like the HDB Block or Facade Assembly is placed on the board, the system physically registers it and updates the virtual game state instantly.",
        "HAZIQ: It was a massive engineering challenge to get the physical pieces reading flawlessly in real-time, but seeing the physical-to-digital synchronization actually working makes the sleepless nights entirely worth it."
    ])

    # SLIDE 10
    add_slide(10, "Summary", [
        "CUE: Press Right Arrow.",
        "HAZIQ: To sum things up smoothly... Kopitalk is a digital bridge built to resolve a physical disconnect.",
        "HAZIQ: It is about bringing people back together through play, no matter what generation they're from.",
        "HAZIQ: Thank you all so much for listening to our presentation."
    ])

    output_path = r"C:\Users\Haziq\Documents\Kopitalk Slides-20260409T153345Z-3-001\Kopitalk Slides\Kopitalk_Presentation_Script.docx"
    doc.save(output_path)
    print(f"Successfully generated DOCX at {output_path}")

if __name__ == "__main__":
    main()
